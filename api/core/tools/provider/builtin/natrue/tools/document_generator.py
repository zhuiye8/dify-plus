"""
文档生成工具实现
"""
from typing import Any, Union, ClassVar, Optional, List, Dict
from docx import Document
import os
import json
import requests
import re
from io import BytesIO

from core.tools.entities.tool_entities import ToolInvokeMessage
from core.tools.tool.builtin_tool import BuiltinTool
from core.file.enums import FileType, FileAttribute
from core.file.file_manager import download, get_attr


class DocumentGeneratorTool(BuiltinTool):
    """
    文档生成工具类
    
    从Word模板生成文档，支持动态数据替换和语言模型增强
    """
    DIFY_API_URL: ClassVar[str] = "https://api.dify.ai/v1"

    def _invoke_dify_llm(self, prompt: str, api_key: str) -> str:
        """调用 Dify 语言模型 API"""
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        data = {
            "messages": [{"role": "user", "content": prompt}],
            "response_mode": "blocking"
        }
        
        try:
            response = requests.post(f"{self.DIFY_API_URL}/chat-messages", headers=headers, json=data)
            response.raise_for_status()
            result = response.json()
            return result.get("answer", "").strip()
        except requests.exceptions.RequestException as e:
            raise Exception(f"语言模型调用失败: {str(e)}")
        except json.JSONDecodeError:
            raise Exception(f"语言模型返回的JSON格式无效: {response.text}")

    def _get_file_content(self, file_obj: Any) -> bytes:
        """
        从文件对象中获取文件内容
        """
        try:
            # 打印文件对象信息用于调试
            print(f"文件对象类型: {type(file_obj)}")
            print(f"文件对象属性: {dir(file_obj)}")
            
            # 如果是 core.file.models.File 类型
            if hasattr(file_obj, 'path') and os.path.exists(file_obj.path):
                with open(file_obj.path, 'rb') as f:
                    return f.read()
            
            # 尝试使用 download 函数
            file_content = download(file_obj)
            if file_content:
                return file_content
                
            raise ValueError(f"不支持的文件对象类型: {type(file_obj)}")
            
        except Exception as e:
            raise ValueError(f"读取文件内容失败: {str(e)}")

    def _extract_placeholders(self, text: str) -> List[str]:
        """从文本中提取 {{field}} 格式的占位符"""
        pattern = r"\{\{([^{}]+)\}\}"
        matches = re.findall(pattern, text)
        return list(set(matches))

    def _get_nested_value(self, data: Any, field_path: str) -> Any:
        """
        获取嵌套数据中的值
        例如: education[0].degree -> data['education'][0]['degree']
        """
        try:
            # 解析字段路径
            parts = re.findall(r'([^\[\].]+)(?:\[(\d+)\])?\.?', field_path)
            value = data
            
            for key, index in parts:
                if not value:
                    return None
                    
                value = value.get(key)
                if index:  # 如果有数组索引
                    try:
                        value = value[int(index)]
                    except (IndexError, TypeError):
                        return None
                        
            return value
        except Exception:
            return None

    def analyze_template(self, doc: Document, data: Dict[str, Any], api_key: str) -> Dict[str, Any]:
        """分析文档模板中的占位符字段并匹配数据字段"""
        all_text = []
        for para in doc.paragraphs:
            all_text.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        all_text.append(para.text)
        
        template_text = "\n".join(all_text)
        extracted_placeholders = self._extract_placeholders(template_text)
        
        # 准备替换数据
        replacement_data = {}
        for field in extracted_placeholders:
            value = self._get_nested_value(data, field)
            if value is not None:
                replacement_data[field] = value
        
        return replacement_data

    def polish_field(self, field_name: str, field_value: str, api_key: str) -> str:
        """使用语言模型润色字段值"""
        prompt = f"请润色以下文本内容（这是'{field_name}'字段的内容），使其更专业、流畅自然，但保持原始含义：\n\n{field_value}"
        return self._invoke_dify_llm(prompt, api_key)

    def _invoke(
        self,
        user_id: str,
        tool_parameters: dict[str, Any],
    ) -> Union[ToolInvokeMessage, list[ToolInvokeMessage]]:
        """
        执行文档生成逻辑
        """
        # 获取凭证和参数
        dify_api_key = self.runtime.credentials.get("dify_api_key")
        if not dify_api_key:
            raise Exception("Dify API Key 未提供")

        template_file = tool_parameters.get("template")
        data_json = tool_parameters.get("data")
        polish_fields_str = tool_parameters.get("polish_fields", "")
        polish_fields = [field.strip() for field in polish_fields_str.split(",")] if polish_fields_str else []

        # 校验输入
        if not template_file:
            raise Exception("模板文件未提供")
        if not data_json:
            raise Exception("数据未提供")

        # 解析 JSON 数据
        try:
            data = json.loads(data_json) if isinstance(data_json, str) else data_json
        except json.JSONDecodeError:
            raise Exception("数据格式错误，必须为有效的 JSON")

        try:
            # 获取文件内容
            template_content = self._get_file_content(template_file)
            
            # 使用BytesIO处理文档
            doc = Document(BytesIO(template_content))
        except Exception as e:
            raise Exception(f"读取模板文件失败: {str(e)}")

        # 分析模板中的字段并获取替换数据
        replacement_data = self.analyze_template(doc, data, dify_api_key)

        # 处理需要润色的字段
        if polish_fields:
            for field in polish_fields:
                if field in replacement_data:
                    value = replacement_data[field]
                    if value is not None:
                        replacement_data[field] = self.polish_field(field, str(value), dify_api_key)
        
        # 替换文档中的占位符
        self._replace_placeholders(doc, replacement_data)
        
        # 使用BytesIO保存文档
        output_buffer = BytesIO()
        doc.save(output_buffer)
        output_buffer.seek(0)
        file_content = output_buffer.read()

        # 生成文件名
        output_filename = "generated_document.docx"
        if hasattr(template_file, 'name'):
            base_name = os.path.splitext(template_file.name)[0]
            output_filename = f"{base_name}_generated.docx"

        return [
            self.create_text_message("文档生成成功"),
            self.create_blob_message(
                blob=file_content,
                meta={
                    "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "name": output_filename
                },
                save_as=self.VariableKey.DOCUMENT
            )
        ]

    def _replace_placeholders(self, doc: Document, replacement_data: Dict[str, str]) -> None:
        """替换文档中所有位置的占位符"""
        for para in doc.paragraphs:
            self._replace_text_in_paragraph(para, replacement_data)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_text_in_paragraph(para, replacement_data)
    
    def _replace_text_in_paragraph(self, para, replacement_data: Dict[str, str]) -> None:
        """替换段落中的文本"""
        text = para.text
        new_text = text
        
        for key, value in replacement_data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in new_text:
                new_text = new_text.replace(placeholder, str(value))
        
        if new_text != text:
            para.text = new_text

    def validate_credentials(self, credentials: dict[str, Any], tool_parameters: dict[str, Any]) -> None:
        """验证凭证和参数"""
        dify_api_key = credentials.get("dify_api_key")
        if not dify_api_key:
            raise Exception("Dify API Key 未提供")

        # 检查必要参数是否存在
        if "template" not in tool_parameters:
            raise Exception("模板文件参数缺失")
        if "data" not in tool_parameters:
            raise Exception("数据参数缺失")

        # 验证JSON数据格式
        data_json = tool_parameters.get("data")
        if isinstance(data_json, str):
            try:
                json.loads(data_json)
            except json.JSONDecodeError:
                raise Exception("数据格式错误，必须为有效的 JSON")

        # 测试语言模型调用
        try:
            self._invoke_dify_llm("测试提示", dify_api_key)
        except Exception as e:
            raise Exception(f"语言模型调用测试失败: {str(e)}") 