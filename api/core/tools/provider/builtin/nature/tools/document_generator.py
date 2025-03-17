"""
文档生成工具实现
"""
from typing import Any, Union, Optional, List, Dict, Set
from docx import Document
import os
import json
import re
from io import BytesIO

from core.tools.entities.tool_entities import ToolInvokeMessage
from core.tools.tool.builtin_tool import BuiltinTool
from core.file.file_manager import download

class DocumentGeneratorTool(BuiltinTool):
    """
    文档生成工具类
    
    从Word模板生成文档，支持动态数据替换。如果占位符对应的数据缺失，则删除包含该占位符的段落。
    """

    def _get_file_content(self, file_obj: Any) -> bytes:
        """
        从文件对象中获取文件内容
        """
        try:
            # 如果是 core.file.models.File 类型
            if hasattr(file_obj, 'path') and os.path.exists(file_obj.path):
                with open(file_obj.path, 'rb') as f:
                    return f.read()
            
            # 尝试使用 download 函数
            file_content = download(file_obj)
            if file_content:
                return file_content
                
            raise ValueError(f"不支持的文件对象类型: {type(file_obj)}")
            
        except Exception as e:
            raise ValueError(f"读取文件内容失败: {str(e)}")

    def _extract_placeholders(self, text: str) -> List[str]:
        """从文本中提取 {{field}} 格式的占位符"""
        pattern = r"\{\{([^{}]+)\}\}"
        matches = re.findall(pattern, text)
        return list(set(matches))

    def _get_nested_value(self, data: Any, field_path: str) -> Any:
        """
        获取嵌套数据中的值
        例如: education[0].degree -> data['education'][0]['degree']
        """
        try:
            parts = re.findall(r'([^\[\].]+)(?:\[(\d+)\])?\.?', field_path)
            value = data
            
            for key, index in parts:
                if not value:
                    return None
                    
                if isinstance(value, dict):
                    value = value.get(key)
                elif isinstance(value, list):
                    try:
                        value = value[int(key)]
                    except (IndexError, ValueError):
                        return None
                else:
                    return None
                    
                if index:  # 如果有数组索引
                    try:
                        value = value[int(index)]
                    except (IndexError, TypeError):
                        return None
                        
            return value
        except Exception:
            return None

    def _get_array_index_from_field(self, field: str) -> Optional[int]:
        """从字段路径中提取数组索引"""
        match = re.search(r'\[(\d+)\]', field)
        if match:
            return int(match.group(1))
        return None

    def _get_array_name_from_field(self, field: str) -> Optional[str]:
        """从字段路径中提取数组名称"""
        match = re.match(r'([^\[\]]+)(?:\[\d+\])?', field)
        if match:
            return match.group(1)
        return None

    def analyze_template(self, doc: Document, data: Dict[str, Any]) -> Dict[str, Any]:
        """分析文档模板中的占位符字段并匹配数据字段"""
        all_text = []
        for para in doc.paragraphs:
            all_text.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        all_text.append(para.text)
        
        template_text = "\n".join(all_text)
        extracted_placeholders = self._extract_placeholders(template_text)
        print("提取到的占位符:", extracted_placeholders)
        
        # 准备替换数据
        replacement_data = {}
        for field in extracted_placeholders:
            value = self._get_nested_value(data, field)
            if value is not None:
                replacement_data[field] = value
                print(f"字段 {field} 的替换值: {value}")
            else:
                replacement_data[field] = None  # 明确设置为 None，便于后续处理
                print(f"字段 {field} 未找到对应的值")
        
        print("最终替换数据:", replacement_data)
        return replacement_data

    def _invoke(
        self,
        user_id: str,
        tool_parameters: dict[str, Any],
    ) -> Union[ToolInvokeMessage, list[ToolInvokeMessage]]:
        """
        执行文档生成逻辑
        """
        template_file = tool_parameters.get("template")
        data_json = tool_parameters.get("data")

        # 校验输入
        if not template_file:
            raise Exception("模板文件未提供")
        if not data_json:
            raise Exception("数据未提供")

        # 解析 JSON 数据
        try:
            data = json.loads(data_json) if isinstance(data_json, str) else data_json
            if isinstance(data, dict):
                print("输入数据:", {k: v for k, v in data.items() if not hasattr(v, '__dict__')})
            elif isinstance(data, list):
                print("输入数据:", data)
                if len(data) > 0:
                    data = data[0]
                else:
                    raise Exception("数据列表为空")
        except json.JSONDecodeError:
            raise Exception("数据格式错误，必须为有效的 JSON")

        try:
            # 获取文件内容
            template_content = self._get_file_content(template_file)
            doc = Document(BytesIO(template_content))
        except Exception as e:
            raise Exception(f"读取模板文件失败: {str(e)}")

        # 分析模板中的字段并获取替换数据
        replacement_data = self.analyze_template(doc, data)
        
        # 替换文档中的占位符
        self._replace_placeholders(doc, replacement_data)
        
        # 使用 BytesIO 保存文档
        output_buffer = BytesIO()
        doc.save(output_buffer)
        output_buffer.seek(0)
        file_content = output_buffer.read()

        # 生成文件名
        output_filename = "generated_document.docx"
        if hasattr(template_file, 'name'):
            base_name = os.path.splitext(template_file.name)[0]
            output_filename = f"{base_name}_generated.docx"

        return [
            self.create_text_message("文档生成成功"),
            self.create_blob_message(
                blob=file_content,
                meta={
                    "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "name": output_filename
                },
                save_as=self.VariableKey.DOCUMENT
            )
        ]

    def _replace_placeholders(self, doc: Document, replacement_data: Dict[str, str]) -> None:
        """替换文档中所有位置的占位符，如果数据缺失则删除对应段落"""
        # 收集所有无效的数组索引
        invalid_indices: Set[int] = set()
        for field in replacement_data.keys():
            if replacement_data[field] is None:
                index = self._get_array_index_from_field(field)
                if index is not None:
                    invalid_indices.add(index)
        
        # 处理段落
        paragraphs_to_remove = []
        for para in doc.paragraphs:
            if not self._replace_text_in_paragraph(para, replacement_data, invalid_indices):
                paragraphs_to_remove.append(para)
        
        # 删除需要移除的段落
        for para in paragraphs_to_remove:
            p = para._element
            p.getparent().remove(p)
        
        # 处理表格
        for table in doc.tables:
            rows_to_remove = []
            for row in table.rows:
                cells_to_remove = []
                for cell in row.cells:
                    paragraphs_to_remove = []
                    for para in cell.paragraphs:
                        if not self._replace_text_in_paragraph(para, replacement_data, invalid_indices):
                            paragraphs_to_remove.append(para)
                    
                    for para in paragraphs_to_remove:
                        p = para._element
                        p.getparent().remove(p)
                    
                    if not cell.paragraphs:
                        cells_to_remove.append(cell)
                
                for cell in cells_to_remove:
                    cell._element.getparent().remove(cell._element)
                
                if not row.cells:
                    rows_to_remove.append(row)
            
            for row in rows_to_remove:
                row._element.getparent().remove(row._element)
    
    def _replace_text_in_paragraph(self, para, replacement_data: Dict[str, str], invalid_indices: Set[int]) -> bool:
        """
        替换段落中的文本
        返回: bool - 如果段落应该保留返回 True，否则返回 False
        """
        text = para.text
        new_text = text
        
        # 检查段落中是否包含任何占位符
        has_placeholders = False
        for key in replacement_data.keys():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in new_text:
                has_placeholders = True
                value = replacement_data[key]
                
                # 检查是否包含无效的数组索引
                index = self._get_array_index_from_field(key)
                if index is not None and index in invalid_indices:
                    return False
                
                if value is not None:
                    new_text = new_text.replace(placeholder, str(value))
                else:
                    # 如果任何占位符的值为 None，整个段落都应该被删除
                    return False
        
        # 如果段落中没有占位符，保留原样
        if not has_placeholders:
            return True
            
        # 如果段落中有占位符且所有值都有效，更新文本
        if new_text != text:
            para.text = new_text
        return True

    def validate_credentials(self, credentials: dict[str, Any], tool_parameters: dict[str, Any]) -> None:
        """验证参数"""
        if "template" not in tool_parameters:
            raise Exception("模板文件参数缺失")
        if "data" not in tool_parameters:
            raise Exception("数据参数缺失")

        data_json = tool_parameters.get("data")
        if isinstance(data_json, str):
            try:
                json.loads(data_json)
            except json.JSONDecodeError:
                raise Exception("数据格式错误，必须为有效的 JSON")
